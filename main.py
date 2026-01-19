from docx import Document
import npttf2utf
from tkinter import *
from tkinter import filedialog,  messagebox
import os
from datetime import datetime

LEGACY_FONTS = {
    "preeti",
    "himalayan tt",
    "kantipur",
    "pcs nepali",
}

def is_legacy_font(run):
    # Direct font
    if run.font.name and run.font.name.lower() in LEGACY_FONTS:
        return True
    # Inherited from style
    if run.style and run.style.font.name:
        return run.style.font.name.lower() in LEGACY_FONTS
    return False


def convert_runs(runs, mapper):
    for run in runs:
        if is_legacy_font(run) and run.text.strip():
            run.text = mapper.map_to_unicode(run.text, from_font="Preeti")
            run.font.name = "Nirmala UI"  # Unicode Nepali font


def convert_docx_preserve_everything(input_docx, output_docx, map_file="map.json"):
    doc = Document(input_docx)
    mapper = npttf2utf.FontMapper(map_file)
    # Convert paragraphs
    for para in doc.paragraphs:
        convert_runs(para.runs, mapper)
    # Convert tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    convert_runs(para.runs, mapper)
    # Headers & footers (often missed!)
    for section in doc.sections:
        for para in section.header.paragraphs:
            convert_runs(para.runs, mapper)
        for para in section.footer.paragraphs:
            convert_runs(para.runs, mapper)
    doc.save(output_docx)


# custom warning 
def show_red_warning(title, message):
    win = Toplevel(root)
    win.title(title)
    win.configure(bg="red")
    win.geometry("450x200")
    win.resizable(False, False)

    # Create a Text widget for copyable message
    text_widget = Text(win, bg="red", fg="white", font=("Arial", 12, "bold"), wrap="word")
    text_widget.insert("1.0", message)
    text_widget.config(state="disabled")  # make it read-only
    text_widget.pack(expand=True, fill="both", padx=20, pady=20)
    # Button to close
    Button(win, text="OK", command=win.destroy, bg="white", fg="red", width=10).pack(pady=10)
    win.grab_set()  # modal window


def select_input_file():
    file_path = filedialog.askopenfilename(title="Select Word File", filetypes=[("Word files","*.docx")])
    if file_path:
        input_var.set(file_path)
        show_output_field()

def select_output_file():
    file_path = filedialog.asksaveasfilename(title="Save As", defaultextension=".docx",
                                             filetypes=[("Word files","*.docx")])
    if file_path:
        output_var.set(file_path)

def convert_file():
    input_path = input_var.get()
    output_path = output_var.get()
    if not input_path or not os.path.exists(input_path):
        messagebox.showwarning("Warning", "Please select a valid input file.")
        return
    if not output_path:
        messagebox.showwarning("Warning", "Please specify output file path.")
        return
    try:
        if not output_path.lower().endswith(".docx"):
            show_red_warning("Invalid file", "Please save the file with a .docx extension.")
            return
        convert_docx_preserve_everything(input_path, output_path)
        messagebox.showinfo("Success", f"File converted and saved to:\n{output_path}")
    except Exception as e:
        messagebox.showerror("Error", f"Conversion failed:\n{e}")
    
    
def show_output_field():
    default_output = os.path.join(os.getcwd(), "output", f"output_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx")
    # Match the slash style of input file    
    output_var.set('default_output')    
    if '/' in input_var.get():
        default_output = default_output.replace('\\', '/')
    else:
        default_output = default_output.replace('/', '\\')
    output_var.set(default_output)
    Label(root, text="Output File:").grid(row=1, column=0, padx=20, pady=20, sticky=E)
    Entry(root, textvariable=output_var, width=80).grid(row=1, column=1, padx=20, pady=20, sticky=W) 
    Button(root, text="Browse", command=select_output_file, width=15).grid(row=1, column=2, padx=20)

      
if __name__ == "__main__":
    root = Tk()
    root.title("Nepali Unicode Converter")
    # root.geometry(f"{root.winfo_screenwidth()}x{root.winfo_screenheight()}")  # full page

    root.columnconfigure(0, weight=1)
    root.columnconfigure(1, weight=3)

    input_var = StringVar()
    output_var = StringVar()
    Label(root, text="Input File:").grid(row=0, column=0, padx=20, pady=20, sticky=E)
    Entry(root, textvariable=input_var, width=80).grid(row=0, column=1, padx=20, pady=20, sticky=W)
    Button(root, text="Browse", command=select_input_file, width=15).grid(row=0, column=2, padx=20)

    # Output widgets (always visible)
    Label(root, text="Output File:").grid(row=1, column=0, padx=20, pady=20, sticky=E)
    Entry(root, textvariable=output_var, width=80).grid(row=1, column=1, padx=20, pady=20, sticky=W)
    Button(root, text="Browse", command=select_output_file, width=15).grid(row=1, column=2, padx=20)


    Button(root, text="Convert to Unicode", command=convert_file, bg="#4CAF50", fg="white",
        font=("Arial", 14, "bold"), width=25, height=2).grid(row=2, column=0, columnspan=3, pady=40)

    root.mainloop()
    